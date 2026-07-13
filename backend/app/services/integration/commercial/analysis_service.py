"""Commercial bid analysis queries and Word report export."""

from __future__ import annotations

import json
import re
import tempfile
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from app.services.integration.commercial.export_service import CommercialExportService, _split_company_names
from app.services.integration.commercial.flat_ingest import is_win_status, normalize_purchaser_label
from app.services.integration.commercial.ic_ingest_service import normalize_enterprise_name
from app.services.shared.db.sqlite_client import SqliteClient


@dataclass(frozen=True)
class CommercialAnalysisFilters:
    company_name: str = ""
    purchaser: str = ""
    inquiry_no: str = ""
    winner: str = ""
    amount_min: float | None = None
    amount_max: float | None = None
    participation_min: int | None = None
    only_winners: bool = False
    start_time: str = ""
    end_time: str = ""


class CommercialAnalysisService:
    """Query 商务网数据 and build bid-fund statistics."""

    def __init__(self, client: SqliteClient | None = None) -> None:
        self._client = client or SqliteClient()

    def filter_options(self, commercial_batch_id: str) -> dict[str, list[str]]:
        records = self._load_records(commercial_batch_id)
        return {
            "company_name": self._distinct(records, "company_name"),
            "purchaser": self._distinct(records, "purchaser"),
            "inquiry_no": self._distinct(records, "inquiry_no"),
            "winner": self._distinct_winners(records),
            "source": self._distinct(records, "source"),
        }

    def query_records(
        self,
        commercial_batch_id: str,
        filters: CommercialAnalysisFilters | None = None,
        limit: int = 5000,
    ) -> dict[str, Any]:
        active = filters or CommercialAnalysisFilters()
        records = [r for r in self._load_records(commercial_batch_id) if self._match_filters(r, active)]

        if active.participation_min is not None and int(active.participation_min) > 0:
            min_count = int(active.participation_min)
            pref_summary = self.summarize(records, commercial_batch_id)
            allowed = {
                c["company_norm"]
                for c in pref_summary["company_summary"]
                if int(c["participation_count"]) >= min_count
            }
            records = [
                r
                for r in records
                if normalize_enterprise_name(self._to_text(r.get("company_name"))) in allowed
            ]

        summary = self.summarize(records, commercial_batch_id)
        records_out = records[: max(1, min(int(limit), 10000))]
        return {
            "records": records_out,
            "summary": summary,
            "description": self.render_description(summary),
        }

    def summarize(self, records: list[dict[str, Any]], commercial_batch_id: str = "") -> dict[str, Any]:
        company_inquiries: dict[str, set[str]] = defaultdict(set)
        company_wins: dict[str, set[str]] = defaultdict(set)
        company_amounts: dict[str, float] = defaultdict(float)
        purchaser_amounts: dict[str, float] = defaultdict(float)
        fund_key_amount: dict[tuple[str, str, str], float] = {}
        company_display: dict[str, str] = {}
        purchaser_inquiries: dict[str, set[str]] = defaultdict(set)

        risk_map = self._risk_summary_map(commercial_batch_id) if commercial_batch_id else {}
        for row in records:
            company = self._to_text(row.get("company_name"))
            cn = normalize_enterprise_name(company)
            if not cn:
                continue
            company_display.setdefault(cn, company)
            inquiry = self._to_text(row.get("inquiry_no"))
            purchaser = self._to_text(row.get("purchaser"))
            amount = float(row.get("win_amount") or 0)
            if inquiry:
                company_inquiries[cn].add(inquiry)
                if purchaser:
                    purchaser_inquiries[purchaser].add(inquiry)
            if bool(row.get("is_winner")):
                company_wins[cn].add(inquiry)
                key = (cn, purchaser, inquiry)
                fund_key_amount[key] = max(fund_key_amount.get(key, 0.0), amount)

        for (cn, purchaser, _inquiry), amount in fund_key_amount.items():
            company_amounts[cn] += amount
            if purchaser:
                purchaser_amounts[purchaser] += amount

        company_summary = []
        for cn, inquiries in company_inquiries.items():
            risk = risk_map.get(cn, {})
            company_summary.append(
                {
                    "company_name": company_display.get(cn, cn),
                    "company_norm": cn,
                    "participation_count": len(inquiries),
                    "win_count": len(company_wins.get(cn, set())),
                    "win_amount": round(company_amounts.get(cn, 0.0), 2),
                    "risk_level": risk.get("risk_level", ""),
                    "risk_score": risk.get("total_score", 0),
                    "risk_hit_count": risk.get("hit_count", 0),
                }
            )
        company_summary.sort(key=lambda x: (-float(x["win_amount"]), -int(x["win_count"]), x["company_name"]))

        purchaser_summary = [
            {
                "purchaser": purchaser,
                "inquiry_count": len(purchaser_inquiries.get(purchaser, set())),
                "win_amount": round(amount, 2),
            }
            for purchaser, amount in purchaser_amounts.items()
        ]
        purchaser_summary.sort(key=lambda x: (-float(x["win_amount"]), x["purchaser"]))

        fund_links = []
        for row in records:
            if not bool(row.get("is_winner")):
                continue
            cn = normalize_enterprise_name(self._to_text(row.get("company_name")))
            risk = risk_map.get(cn, {})
            fund_links.append(
                {
                    "company_name": row.get("company_name", ""),
                    "purchaser": row.get("purchaser", ""),
                    "inquiry_no": row.get("inquiry_no", ""),
                    "winner": row.get("winner", ""),
                    "win_amount": row.get("win_amount", 0),
                    "risk_level": risk.get("risk_level", ""),
                    "risk_score": risk.get("total_score", 0),
                    "source": row.get("source", ""),
                }
            )
        fund_links.sort(key=lambda x: (-float(x["win_amount"] or 0), x["company_name"], x["inquiry_no"]))

        inquiry_count = len({self._to_text(r.get("inquiry_no")) for r in records if self._to_text(r.get("inquiry_no"))})
        return {
            "record_count": len(records),
            "inquiry_count": inquiry_count,
            "company_count": len(company_summary),
            "winner_company_count": len([c for c in company_summary if int(c["win_count"]) > 0]),
            "total_win_amount": round(sum(company_amounts.values()), 2),
            "company_summary": company_summary,
            "purchaser_summary": purchaser_summary,
            "fund_links": fund_links,
            "top_company_amounts": [[r["company_name"], r["win_amount"]] for r in company_summary[:10]],
            "top_purchaser_amounts": [[r["purchaser"], r["win_amount"]] for r in purchaser_summary[:10]],
        }

    def render_description(self, summary: dict[str, Any]) -> str:
        top_company = (summary.get("company_summary") or [{}])[0] if summary.get("company_summary") else {}
        top_purchaser = (summary.get("purchaser_summary") or [{}])[0] if summary.get("purchaser_summary") else {}
        lines = [
            f"本次商务网分析覆盖 {summary.get('inquiry_count', 0)} 个询价单、{summary.get('company_count', 0)} 家参与企业，识别中标企业 {summary.get('winner_company_count', 0)} 家。",
            f"统计口径下中标金额合计 {float(summary.get('total_win_amount') or 0):,.2f} 元，资金关联明细 {len(summary.get('fund_links') or [])} 条。",
        ]
        if top_company:
            lines.append(
                f"按企业中标金额排序，排名第一的是 {top_company.get('company_name', '')}，中标 {top_company.get('win_count', 0)} 次，金额 {float(top_company.get('win_amount') or 0):,.2f} 元。"
            )
        if top_purchaser:
            lines.append(
                f"按采购单位统计，金额最高的是 {top_purchaser.get('purchaser', '')}，关联中标金额 {float(top_purchaser.get('win_amount') or 0):,.2f} 元。"
            )
        return "\n".join(lines)

    def export_statistics_report(self, commercial_batch_id: str, output_path: str) -> str:
        out_file = Path(output_path)
        if out_file.suffix.lower() != ".docx":
            out_file = out_file.with_suffix(".docx")
        out_file.parent.mkdir(parents=True, exist_ok=True)

        result = self.query_records(commercial_batch_id, CommercialAnalysisFilters(), limit=10000)
        summary = result["summary"]

        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            self._export_simple_docx(out_file, commercial_batch_id, result)
            return str(out_file)

        document = Document()
        document.add_heading("商务网统计分析报告", level=0)
        document.add_paragraph(f"商务网批次：{commercial_batch_id}")
        document.add_paragraph(result["description"])

        document.add_heading("一、核心统计", level=1)
        kpi_table = document.add_table(rows=1, cols=2)
        kpi_table.style = "Table Grid"
        kpi_table.rows[0].cells[0].text = "指标"
        kpi_table.rows[0].cells[1].text = "数值"
        for name, value in [
            ("询价单数量", summary["inquiry_count"]),
            ("参与企业数", summary["company_count"]),
            ("中标企业数", summary["winner_company_count"]),
            ("中标金额合计", f"{float(summary['total_win_amount']):,.2f}"),
        ]:
            cells = kpi_table.add_row().cells
            cells[0].text = str(name)
            cells[1].text = str(value)

        with tempfile.TemporaryDirectory() as tmp:
            chart_paths = self._build_report_charts(summary, Path(tmp))
            if chart_paths:
                document.add_heading("二、统计图表", level=1)
                for title, path in chart_paths:
                    document.add_paragraph(title)
                    document.add_picture(str(path), width=Inches(6.2))

        document.add_heading("三、企业中标统计（Top 20）", level=1)
        self._add_table(
            document,
            ["企业", "参标次数", "中标次数", "中标金额", "风险等级", "风险分"],
            [
                [
                    r["company_name"],
                    r["participation_count"],
                    r["win_count"],
                    f"{float(r['win_amount']):,.2f}",
                    r.get("risk_level", ""),
                    r.get("risk_score", 0),
                ]
                for r in summary["company_summary"][:20]
            ],
        )

        document.add_heading("四、中标资金关联（Top 30）", level=1)
        self._add_table(
            document,
            ["企业", "采购单位", "询价单号", "中标金额", "风险等级", "数据来源"],
            [
                [
                    r["company_name"],
                    r["purchaser"],
                    r["inquiry_no"],
                    f"{float(r['win_amount'] or 0):,.2f}",
                    r.get("risk_level", ""),
                    r.get("source", ""),
                ]
                for r in summary["fund_links"][:30]
            ],
        )

        document.save(out_file)
        return str(out_file)

    def _export_simple_docx(
        self,
        out_file: Path,
        commercial_batch_id: str,
        result: dict[str, Any],
    ) -> None:
        summary = result["summary"]
        body: list[str] = [
            self._docx_paragraph("商务网统计分析报告", style="Title"),
            self._docx_paragraph(f"商务网批次：{commercial_batch_id}"),
            self._docx_paragraph(str(result.get("description", ""))),
            self._docx_paragraph("一、核心统计", style="Heading1"),
            self._docx_table(
                ["指标", "数值"],
                [
                    ["询价单数量", summary["inquiry_count"]],
                    ["参与企业数", summary["company_count"]],
                    ["中标企业数", summary["winner_company_count"]],
                    ["中标金额合计", f"{float(summary['total_win_amount']):,.2f}"],
                ],
            ),
            self._docx_paragraph("二、图表展示", style="Heading1"),
            self._docx_paragraph("企业中标金额 Top 10"),
            self._docx_bar_table(summary.get("top_company_amounts") or []),
            self._docx_paragraph("采购单位关联中标金额 Top 10"),
            self._docx_bar_table(summary.get("top_purchaser_amounts") or []),
            self._docx_paragraph("三、企业中标统计（Top 20）", style="Heading1"),
            self._docx_table(
                ["企业", "参标次数", "中标次数", "中标金额", "风险等级", "风险分"],
                [
                    [
                        r["company_name"],
                        r["participation_count"],
                        r["win_count"],
                        f"{float(r['win_amount']):,.2f}",
                        r.get("risk_level", ""),
                        r.get("risk_score", 0),
                    ]
                    for r in summary["company_summary"][:20]
                ],
            ),
            self._docx_paragraph("四、中标资金关联（Top 30）", style="Heading1"),
            self._docx_table(
                ["企业", "采购单位", "询价单号", "中标金额", "风险等级", "数据来源"],
                [
                    [
                        r["company_name"],
                        r["purchaser"],
                        r["inquiry_no"],
                        f"{float(r['win_amount'] or 0):,.2f}",
                        r.get("risk_level", ""),
                        r.get("source", ""),
                    ]
                    for r in summary["fund_links"][:30]
                ],
            ),
        ]
        document_xml = self._docx_document_xml("".join(body))
        with zipfile.ZipFile(out_file, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", self._docx_content_types())
            zf.writestr("_rels/.rels", self._docx_root_rels())
            zf.writestr("word/document.xml", document_xml)

    def _load_records(self, commercial_batch_id: str) -> list[dict[str, Any]]:
        export = CommercialExportService(self._client)
        rows = self._fill_forward(export._load_commercial_rows(commercial_batch_id))
        records: list[dict[str, Any]] = []
        seen: set[tuple[str, str, str, str, float]] = set()
        for row in rows:
            companies = _split_company_names(self._to_text(row.get("公司名称")))
            if not companies:
                continue
            winners = _split_company_names(self._to_text(row.get("中标供应商")))
            winner_norms = {normalize_enterprise_name(w) for w in winners if normalize_enterprise_name(w)}
            status = self._to_text(row.get("状态"))
            base_amount = self._safe_amount(row.get("中标金额(元)", ""))
            for company in companies:
                cn = normalize_enterprise_name(company)
                is_winner_by_supplier = bool(cn and any(self._is_same_company(cn, wn) for wn in winner_norms))
                is_winner = is_win_status(status) or is_winner_by_supplier
                amount = base_amount
                inquiry_no = self._to_text(row.get("询价单号"))
                purchaser = normalize_purchaser_label(self._to_text(row.get("采购单位")))
                key = (cn, inquiry_no, purchaser, self._to_text(row.get("序号")), amount if is_winner else 0.0)
                if key in seen:
                    continue
                seen.add(key)
                records.append(
                    {
                        "source": self._to_text(row.get("数据来源")),
                        "inquiry_no": inquiry_no,
                        "purchaser": purchaser,
                        "company_name": company,
                        "winner": "、".join(winners),
                        "is_winner": is_winner,
                        "win_amount": round(amount if is_winner else 0.0, 2),
                        "item_name": self._to_text(row.get("物资编码/来源采购申请代码--物资描述")),
                        "quote_price": self._to_text(row.get("含税单价(元)")),
                        "quantity": self._to_text(row.get("数量")),
                        "remark": self._to_text(row.get("备注")),
                        "inquiry_time": self._extract_inquiry_time(row),
                        "bid_status": status,
                    }
                )
        return records

    def _risk_summary_map(self, commercial_batch_id: str) -> dict[str, dict[str, Any]]:
        rows = self._client.query_all(
            """
            SELECT enterprise_name, total_score, hit_count, risk_level
            FROM ana_risk_summary
            WHERE import_batch_id=?;
            """,
            (commercial_batch_id,),
        )
        out: dict[str, dict[str, Any]] = {}
        for name, score, hits, level in rows:
            norm = normalize_enterprise_name(str(name or ""))
            if norm:
                out[norm] = {"total_score": score, "hit_count": hits, "risk_level": level}
        return out

    def _match_filters(self, row: dict[str, Any], filters: CommercialAnalysisFilters) -> bool:
        if filters.company_name and filters.company_name not in self._to_text(row.get("company_name")):
            return False
        if filters.purchaser and filters.purchaser not in self._to_text(row.get("purchaser")):
            return False
        if filters.inquiry_no and filters.inquiry_no not in self._to_text(row.get("inquiry_no")):
            return False
        if filters.winner and filters.winner not in self._to_text(row.get("winner")):
            return False
        if filters.only_winners and not bool(row.get("is_winner")):
            return False
        amount = float(row.get("win_amount") or 0)
        if filters.amount_min is not None and amount < float(filters.amount_min):
            return False
        if filters.amount_max is not None and amount > float(filters.amount_max):
            return False
        inquiry_time = self._parse_time(row.get("inquiry_time"))
        if filters.start_time or filters.end_time:
            if inquiry_time is None:
                return False
            if filters.start_time:
                start = self._parse_time(filters.start_time)
                if start and inquiry_time < start:
                    return False
            if filters.end_time:
                end = self._parse_time(filters.end_time)
                if end and inquiry_time > end:
                    return False
        return True

    @staticmethod
    def _parse_time(value: Any) -> datetime | None:
        text = CommercialAnalysisService._to_text(value)
        if not text:
            return None
        normalized = text.replace("/", "-").replace("T", " ")
        for fmt in (
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%d %H:%M",
            "%Y-%m-%d",
        ):
            try:
                return datetime.strptime(normalized, fmt)
            except ValueError:
                continue
        try:
            return datetime.fromisoformat(normalized[:19])
        except ValueError:
            return None

    @staticmethod
    def _extract_inquiry_time(row: dict[str, str]) -> str:
        for key in ("提交人/时间", "报价截止时间", "公布中标信息", "要求到货日期"):
            text = CommercialAnalysisService._to_text(row.get(key))
            if not text:
                continue
            match = re.search(r"(\d{4}[-/]\d{1,2}[-/]\d{1,2}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?)", text)
            if match:
                return match.group(1).replace("/", "-")
        return ""

    @staticmethod
    def _fill_forward(rows: list[dict[str, str]]) -> list[dict[str, str]]:
        out: list[dict[str, str]] = []
        carry: dict[str, str] = {
            "数据来源": "",
            "询价单号": "",
            "公司名称": "",
            "采购单位": "",
            "提交人/时间": "",
            "报价截止时间": "",
        }
        for row in rows:
            item = dict(row)
            for key in carry:
                value = (item.get(key) or "").strip()
                if value:
                    carry[key] = value
                else:
                    item[key] = carry[key]
            out.append(item)
        return out

    @staticmethod
    def _distinct(records: list[dict[str, Any]], key: str) -> list[str]:
        values = sorted({str(r.get(key) or "").strip() for r in records if str(r.get(key) or "").strip()})
        if key == "company_name":
            return values
        return values[:1000]

    @staticmethod
    def _distinct_winners(records: list[dict[str, Any]]) -> list[str]:
        values: set[str] = set()
        for record in records:
            for winner in _split_company_names(str(record.get("winner") or "")):
                if winner and not CommercialAnalysisService._looks_like_datetime_winner(winner):
                    values.add(winner)
        return sorted(values)[:1000]

    @staticmethod
    def _looks_like_datetime_winner(text: str) -> bool:
        value = (text or "").strip()
        if not value:
            return False
        return bool(re.match(r"^\d{4}-\d{2}-\d{2}(?:\s+\d{2}:\d{2}:\d{2})?$", value))

    @staticmethod
    def _to_text(value: Any) -> str:
        return "" if value is None else str(value).strip()

    @staticmethod
    def _safe_amount(value: Any) -> float:
        text = str(value or "").replace(",", "").strip()
        if not text:
            return 0.0
        try:
            return float(text)
        except Exception:
            return 0.0

    @staticmethod
    def _is_same_company(company_norm: str, winner_norm: str) -> bool:
        return bool(company_norm and winner_norm and (company_norm == winner_norm or company_norm in winner_norm or winner_norm in company_norm))

    @staticmethod
    def _add_table(document: Any, headers: list[str], rows: list[list[Any]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = str(header)
        if not rows:
            cells = table.add_row().cells
            cells[0].text = "暂无数据"
            return
        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row[: len(headers)]):
                cells[idx].text = str(value)

    @staticmethod
    def _build_report_charts(summary: dict[str, Any], tmp_dir: Path) -> list[tuple[str, Path]]:
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            return []

        plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False
        chart_paths: list[tuple[str, Path]] = []

        company_top = summary.get("top_company_amounts") or []
        if company_top:
            path = tmp_dir / "company_amounts.png"
            labels = [str(x[0]) for x in company_top][::-1]
            values = [float(x[1] or 0) for x in company_top][::-1]
            fig, ax = plt.subplots(figsize=(8, 4.5))
            ax.barh(labels, values, color="#D99A5B")
            ax.set_title("企业中标金额 Top 10")
            ax.set_xlabel("金额（元）")
            fig.tight_layout()
            fig.savefig(path, dpi=160)
            plt.close(fig)
            chart_paths.append(("企业中标金额 Top 10", path))

        purchaser_top = summary.get("top_purchaser_amounts") or []
        if purchaser_top:
            path = tmp_dir / "purchaser_amounts.png"
            labels = [str(x[0]) for x in purchaser_top][::-1]
            values = [float(x[1] or 0) for x in purchaser_top][::-1]
            fig, ax = plt.subplots(figsize=(8, 4.5))
            ax.barh(labels, values, color="#6B8F71")
            ax.set_title("采购单位关联中标金额 Top 10")
            ax.set_xlabel("金额（元）")
            fig.tight_layout()
            fig.savefig(path, dpi=160)
            plt.close(fig)
            chart_paths.append(("采购单位关联中标金额 Top 10", path))

        return chart_paths

    @staticmethod
    def _docx_content_types() -> str:
        return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

    @staticmethod
    def _docx_root_rels() -> str:
        return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    @staticmethod
    def _docx_document_xml(body: str) -> str:
        return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{body}<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1080" w:bottom="1440" w:left="1080"/></w:sectPr></w:body>
</w:document>"""

    @staticmethod
    def _docx_paragraph(text: str, style: str | None = None) -> str:
        style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        lines = [escape(line) for line in str(text or "").splitlines()] or [""]
        runs = "<w:br/>".join(f"<w:r><w:t>{line}</w:t></w:r>" for line in lines)
        return f"<w:p>{style_xml}{runs}</w:p>"

    @classmethod
    def _docx_table(cls, headers: list[str], rows: list[list[Any]]) -> str:
        def cell(value: Any, shade: str = "") -> str:
            shading = f'<w:tcPr><w:shd w:fill="{shade}"/></w:tcPr>' if shade else ""
            return f"<w:tc>{shading}{cls._docx_paragraph(str(value))}</w:tc>"

        header_xml = "<w:tr>" + "".join(cell(h, "F3E3D2") for h in headers) + "</w:tr>"
        body_rows = rows or [["暂无数据"] + [""] * (len(headers) - 1)]
        rows_xml = "".join(
            "<w:tr>" + "".join(cell(row[idx] if idx < len(row) else "") for idx in range(len(headers))) + "</w:tr>"
            for row in body_rows
        )
        return f"<w:tbl><w:tblPr><w:tblBorders><w:top w:val=\"single\" w:sz=\"4\"/><w:left w:val=\"single\" w:sz=\"4\"/><w:bottom w:val=\"single\" w:sz=\"4\"/><w:right w:val=\"single\" w:sz=\"4\"/><w:insideH w:val=\"single\" w:sz=\"4\"/><w:insideV w:val=\"single\" w:sz=\"4\"/></w:tblBorders></w:tblPr>{header_xml}{rows_xml}</w:tbl>"

    @classmethod
    def _docx_bar_table(cls, pairs: list[list[Any]]) -> str:
        top = pairs[:10]
        max_value = max([float(v or 0) for _, v in top], default=0.0)
        rows: list[list[Any]] = []
        for name, value in top:
            amount = float(value or 0)
            width = 1 if max_value <= 0 else max(1, int(amount / max_value * 24))
            rows.append([name, f"{amount:,.2f}", "█" * width])
        return cls._docx_table(["名称", "金额", "条形图"], rows)


def summary_to_json(summary: dict[str, Any]) -> str:
    return json.dumps(summary, ensure_ascii=False)


__all__ = ["CommercialAnalysisFilters", "CommercialAnalysisService", "summary_to_json"]
